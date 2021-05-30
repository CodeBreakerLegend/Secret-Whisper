from tkinter import *
import tkinter
import tkinter.ttk
import string
import docx
import os
import smtplib

alphalist = string.ascii_lowercase 
Alphalist = string.ascii_uppercase

# smtp_object = smtplib.SMTP("smtp.gmail.com",587)

def encoded(originalmsg,key):
    msg_words = originalmsg.split(" ")
    enc_words = []
    for i in msg_words:
        encw = ""
        for j in i:
            #print(alphalist.find(j)+3,end = " ")
            if j in alphalist:
                enc = alphalist.find(j) + key
            if j in Alphalist:
                enc = Alphalist.find(j) + key
            if enc >= 26:
                enc = enc-26
            #print(alphalist[enc],end="")
            encw = encw + alphalist[enc]
        #print(" ",end="")
        enc_words.append(encw)
    enc_msg = " ".join(enc_words)
    return enc_msg

def decoded(encodedmsg,key):
    enc_msg_list = encodedmsg.split(" ")
    dec_words = []
    for i in enc_msg_list:
        decw = ""
        for j in i:
            #print(alphalist.find(j)+3,end = " ")
            if j in alphalist:
                dec = alphalist.find(j) - key
            if j in Alphalist:
                dec = Alphalist.find(j) - key
            #print(alphalist[dec],end="")
            decw = decw + alphalist[dec]
        #print(" ",end="")
        dec_words.append(decw)
    dec_msg = " ".join(dec_words)
    return dec_msg


def save_info_encode():
    message_info = message.get()
    password_info = int(key.get())
    secret = encoded(message_info,password_info)

    doc = docx.Document()
    parag = doc.add_paragraph(secret)
    doc.save("secret.docx")

    # file = open("secret.txt",'w')
    # file.write(secret)
    # file.close
    print("Secret file successfully generated")

    message_textbox.delete(0, END)
    key_textbox.delete(0, END)
    os.system("start secret.docx")

def save_info_decode():
    message_info = message.get()
    password_info = int(key.get())
    original = decoded(message_info,password_info)

    doc = docx.Document()
    parag = doc.add_paragraph(original)
    doc.save("original.docx")
    
    # file = open("original.docx",'w')
    # file.write(original)
    # file.close
    print("Secret file successfully decrypted")

    message_textbox.delete(0, END)
    key_textbox.delete(0, END)

    os.system("start original.docx")
    
def loginApp():
    global smtp_object
    smtp_object = smtplib.SMTP("smtp.gmail.com",587)
    smtp_object.ehlo()
    smtp_object.starttls()
    email = username.get()
    passwrd = password.get()
    smtp_object.login(email,passwrd)
    username_textbox.delete(0, END)
    password_textbox.delete(0, END)
    first_frame.grid_forget()
    third_frame.grid_forget()
    second_frame.grid(column=0, row=0, padx=50, pady=5, sticky=(tkinter.W, tkinter.N, tkinter.E))

def send_email():
    from_address = from_add.get()
    to_address = send_to.get()
    subject = mail_sub.get()
    mesage = encoded(message.get(), int(passkey.get())) 
    msg = "Subject: " + subject + "\n" + mesage
    smtp_object.sendmail(from_address, to_address, msg)
    send_to_textbox.delete(0, END)
    mail_subject_textbox.delete(0, END)
    message_textbox.delete(0, END)
    passkey_textbox.delete(0, END)



def logout():
    smtp_object.quit()
    third_frame.grid_forget()
    second_frame.grid_forget()
    first_frame.grid(column=0, row=0, padx=50, pady=5, sticky=(tkinter.W, tkinter.N, tkinter.E))

def create_widgets_in_first_frame():
    # Create the label for the frame
    first_window_label = tkinter.ttk.Label(first_frame, text='Login Page')
    first_window_label.grid(column=0, row=0, pady=10, padx=10, sticky=(tkinter.N))
    
    #Create labels for text boxes
    username_label = tkinter.ttk.Label(first_frame, text='Username')
    username_label.grid(column=0, row=1, pady=10, padx=10, sticky=(tkinter.N))
    password_label = tkinter.ttk.Label(first_frame, text='Password')
    password_label.grid(column=0, row=2, pady=10, padx=10, sticky=(tkinter.N))
    
    #Create text boxes
#     username = StringVar()
#     password = StringVar()
    global username_textbox,password_textbox
    username_textbox = tkinter.ttk.Entry(first_frame, textvariable = username)
    username_textbox.grid(column=1, row=1, pady=10, padx=10, sticky=(tkinter.N))
    password_textbox = tkinter.ttk.Entry(first_frame, textvariable = password, show="*")
    password_textbox.grid(column=1, row=2, pady=10, padx=10, sticky=(tkinter.N))
    # Create the button for the frame
    first_window_quit_button = tkinter.Button(first_frame, text = "Submit", command = loginApp )
    first_window_quit_button.grid(column=0, row=3, pady=10, sticky=(tkinter.N))
    
def create_widgets_in_second_frame():
    # Create the label for the frame
    second_window_label = tkinter.ttk.Label(second_frame, text='Encode or Decode')
    second_window_label.grid(column=0, row=0, pady=10, padx=10, sticky=(tkinter.N))
    
    #Create labels for text boxes
    message_label = tkinter.ttk.Label(second_frame, text='Message')
    message_label.grid(column=0, row=1, pady=10, padx=10, sticky=(tkinter.N))
    key_label = tkinter.ttk.Label(second_frame, text='Passkey')
    key_label.grid(column=0, row=2, pady=10, padx=10, sticky=(tkinter.N))
    
    #Create text boxes
#     message = StringVar()
#     passkey = StringVar()
    global message_textbox,key_textbox
    message_textbox = tkinter.ttk.Entry(second_frame, textvariable = message)
    message_textbox.grid(column=1, row=1, pady=10, padx=10, sticky=(tkinter.N))
    key_textbox = tkinter.ttk.Entry(second_frame, textvariable = key, show="*")
    key_textbox.grid(column=1, row=2, pady=10, padx=10, sticky=(tkinter.N))
    
    # Create the button for the frame
    second_window_encode_button = tkinter.Button(second_frame, text = "Encode", command = save_info_encode)
    second_window_encode_button.grid(column=0, row=3, pady=10, sticky=(tkinter.N))
    second_window_decode_button = tkinter.Button(second_frame, text = "Decode", command = save_info_decode)
    second_window_decode_button.grid(column=1, row=3, pady=10, sticky=(tkinter.N))

    second_window_logout_button = tkinter.Button(second_frame, text = "Logout", command = logout)
    second_window_logout_button.grid(column=0, row=4, pady=10, sticky=(tkinter.N))
    second_window_mail_button = tkinter.Button(second_frame, text = "Mail", command = call_third_frame_on_top)
    second_window_mail_button.grid(column=1, row=4, pady=10, sticky=(tkinter.N))

def create_widgets_in_third_frame():
    # Create the label for the frame
    third_window_label = tkinter.ttk.Label(third_frame, text='Mail')
    third_window_label.grid(column=0, row=0, pady=10, padx=10, sticky=(tkinter.N))
    
    #Create labels for text boxes
    from_label = tkinter.ttk.Label(third_frame, text='From')
    from_label.grid(column=0, row=1, pady=10, padx=10, sticky=(tkinter.N))
    send_to_label = tkinter.ttk.Label(third_frame, text='Send To')
    send_to_label.grid(column=0, row=2, pady=10, padx=10, sticky=(tkinter.N))
    subject_label = tkinter.ttk.Label(third_frame, text='Subject')
    subject_label.grid(column=0, row=3, pady=10, padx=10, sticky=(tkinter.N))
    message_label = tkinter.ttk.Label(third_frame, text='Message')
    message_label.grid(column=0, row=4, pady=10, padx=10, sticky=(tkinter.N))
    passkey_label = tkinter.ttk.Label(third_frame, text='Passkey')
    passkey_label.grid(column=0, row=5, pady=10, padx=10, sticky=(tkinter.N))
    
    #Create text boxes
#     from_add = StringVar() 
#     send_to = StringVar()
#     message = StringVar()
#     passkey = StringVar()

    global send_to_textbox,mail_subject_textbox,message_textbox,passkey_textbox
    from_textbox = tkinter.ttk.Entry(third_frame, textvariable = from_add)
    from_textbox.grid(column=1, row=1, pady=10, padx=10, sticky=(tkinter.N))
    send_to_textbox = tkinter.ttk.Entry(third_frame, textvariable = send_to)
    send_to_textbox.grid(column=1, row=2, pady=10, padx=10, sticky=(tkinter.N))
    mail_subject_textbox = tkinter.ttk.Entry(third_frame, textvariable = mail_sub)
    mail_subject_textbox.grid(column=1, row=3, pady=10, padx=10, sticky=(tkinter.N))
    message_textbox = tkinter.ttk.Entry(third_frame, textvariable = message)
    message_textbox.grid(column=1, row=4, pady=10, padx=10, sticky=(tkinter.N))
    passkey_textbox = tkinter.ttk.Entry(third_frame, textvariable = passkey, show="*")
    passkey_textbox.grid(column=1, row=5, pady=10, padx=10, sticky=(tkinter.N))
    
    # Create the button for the frame
    third_window_send_button = tkinter.Button(third_frame, text = "Send", command = send_email)
    third_window_send_button.grid(column=0, row=6, pady=10, sticky=(tkinter.N))
    third_window_logout_button = tkinter.Button(third_frame, text = "Logout", command = logout)
    third_window_logout_button.grid(column=1, row=6, pady=10, sticky=(tkinter.N))
    
    third_window_back_button = tkinter.Button(third_frame, text = "Back", command = call_second_frame_on_top)
    third_window_back_button.grid(column=0, row=7, pady=10, sticky=(tkinter.N))
    third_window_quit_button = tkinter.Button(third_frame, text = "Quit", command = quit_program)
    third_window_quit_button.grid(column=1, row=7, pady=10, sticky=(tkinter.N))

def call_first_frame_on_top():
    # This function can be called only from the second window.
    # Hide the second window and show the first window.
    second_frame.grid_forget()
    first_frame.grid(column=0, row=0, padx=50, pady=5, sticky=(tkinter.W, tkinter.N, tkinter.E))

def call_second_frame_on_top():
    # This function can be called from the first and third windows.
    # Hide the first and third windows and show the second window.
    first_frame.grid_forget()
    third_frame.grid_forget()
    second_frame.grid(column=0, row=0, padx=50, pady=5, sticky=(tkinter.W, tkinter.N, tkinter.E))

def call_third_frame_on_top():
    # This function can only be called from the second window.
    # Hide the second window and show the third window.
    second_frame.grid_forget()
    third_frame.grid(column=0, row=0, padx=50, pady=5, sticky=(tkinter.W, tkinter.N, tkinter.E))

def quit_program():
    smtp_object.quit()
    root_window.destroy()

###############################
# Main program starts here :) #
###############################

# Create the root GUI window.
root_window = tkinter.Tk()
root_window.geometry("400x400")
root_window.title("Secret Whisper")
# Define window size
# window_width = 500
# window_heigth = 500

username = StringVar()
password = StringVar()

message = StringVar()
key = StringVar()

from_add = StringVar() 
send_to = StringVar()
mail_sub= StringVar()
message = StringVar()
passkey = StringVar()


# Create frames inside the root window to hold other GUI elements. All frames must be created in the main program, otherwise they are not accessible in functions. 
first_frame=tkinter.ttk.Frame(root_window, width=500, height=500)
# first_frame['borderwidth'] = 2
# first_frame['relief'] = 'sunken'
first_frame.grid(column=0, row=0, padx=100, pady=50, sticky=(tkinter.W, tkinter.N, tkinter.E))

second_frame=tkinter.ttk.Frame(root_window, width=750, height=750)
# second_frame['borderwidth'] = 2
# second_frame['relief'] = 'sunken'
second_frame.grid(column=0, row=0, padx=100, pady=50, sticky=(tkinter.W, tkinter.N, tkinter.E))

third_frame=tkinter.ttk.Frame(root_window, width=1000, height=1000)
# third_frame['borderwidth'] = 2
# third_frame['relief'] = 'sunken'
third_frame.grid(column=0, row=0, padx=100, pady=50, sticky=(tkinter.W, tkinter.N, tkinter.E))

# Create all widgets to all frames
create_widgets_in_third_frame()
create_widgets_in_second_frame()
create_widgets_in_first_frame()

# Hide all frames in reverse order, but leave first frame visible (unhidden).
third_frame.grid_forget()
second_frame.grid_forget()

# Start tkinter event - loop
root_window.mainloop()
